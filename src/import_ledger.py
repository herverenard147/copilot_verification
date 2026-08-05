"""Import d'ecritures/bilan depuis un fichier externe (Excel, CSV, Word) --
la piece qui rend le bilan (src/bilan.py) exploitable au-dela des reçus
d'achat : capital, immobilisations, ventes... rien qu'un reçu ne peut fournir.

Colonnes reconnues (insensible a la casse/aux accents, plusieurs alias
acceptes) : Compte (requis), Libelle (optionnel), Debit, Credit.

Jamais de crash sur un fichier mal forme : une ligne invalide est signalee
(numero + raison) et ignoree, le reste du fichier est quand meme importe --
meme philosophie que le reste de l'app (jamais de 500, erreurs structurees).
"""
import io
import unicodedata

import pandas as pd

_COLUMN_ALIASES = {
    "account": {"compte", "account", "numero de compte", "n compte", "code compte", "no compte"},
    "label": {"libelle", "label", "intitule", "designation"},
    "debit": {"debit", "dt"},
    "credit": {"credit", "ct"},
}


def _normalize(text):
    folded = unicodedata.normalize("NFKD", str(text)).encode("ascii", "ignore").decode()
    return folded.strip().lower()


def _match_columns(columns):
    """{nom_colonne_reel: role} pour les colonnes reconnues."""
    mapping = {}
    for col in columns:
        norm = _normalize(col)
        for role, aliases in _COLUMN_ALIASES.items():
            if norm in aliases and role not in mapping.values():
                mapping[col] = role
                break
    return mapping


def _dataframe_from_docx(raw):
    from docx import Document
    doc = Document(io.BytesIO(raw))
    if not doc.tables:
        raise ValueError("Aucun tableau trouvé dans le document Word.")
    table = doc.tables[0]   # le premier tableau du document (le plus courant)
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if len(rows) < 2:
        raise ValueError("Le tableau ne contient pas de lignes de données.")
    header, *data = rows
    return pd.DataFrame(data, columns=header)


def _dataframe_from_bytes(filename, raw):
    """DataFrame brut (colonnes en texte) selon l'extension. ValueError avec
    un message humain si le format est inconnu ou le fichier illisible."""
    name = (filename or "").lower()
    try:
        if name.endswith((".xlsx", ".xlsm")):
            return pd.read_excel(io.BytesIO(raw), dtype=str)
        if name.endswith(".csv"):
            return pd.read_csv(io.BytesIO(raw), dtype=str, sep=None, engine="python")
        if name.endswith(".docx"):
            return _dataframe_from_docx(raw)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Fichier illisible : {type(exc).__name__}") from exc
    raise ValueError("Format non reconnu (attendu : .xlsx, .csv ou .docx).")


def _to_amount(value):
    """Texte -> float, tolerant (virgule decimale, espaces milliers, vide/NaN -> 0)."""
    if value is None:
        return 0.0
    s = str(value).strip().replace(" ", "").replace(" ", "")
    if not s or s.lower() == "nan":
        return 0.0
    s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        raise ValueError(f"montant invalide : {value!r}")


def parse_ledger_file(filename, raw):
    """Parse un fichier (.xlsx/.csv/.docx) en lignes d'ecriture.

    Renvoie (rows, errors) :
    - rows : liste de {account, label, debit, credit} valides.
    - errors : liste de {row, reason} pour les lignes ignorees (compte
      manquant, montant illisible, ligne vide...).

    Ne leve QUE si le fichier lui-meme est illisible ou si aucune colonne
    "Compte" n'est reconnue -- rien a importer du tout. Une ligne
    individuelle mal formee n'arrete jamais l'import des autres."""
    df = _dataframe_from_bytes(filename, raw)
    if df.empty:
        raise ValueError("Le fichier ne contient aucune ligne.")

    colmap = _match_columns(df.columns)
    if "account" not in colmap.values():
        raise ValueError(
            "Aucune colonne « Compte » reconnue. Colonnes trouvées : "
            + ", ".join(str(c) for c in df.columns)
        )
    inv = {role: col for col, role in colmap.items()}   # role -> nom de colonne reel

    rows, errors = [], []
    for i, raw_row in df.iterrows():
        line_no = i + 2   # ligne 1 = en-tete, index pandas 0-based -> ligne fichier
        account_raw = raw_row.get(inv["account"])
        account = "" if account_raw is None else str(account_raw).strip()
        if not account or account.lower() == "nan":
            errors.append({"row": line_no, "reason": "compte manquant"})
            continue
        try:
            debit = _to_amount(raw_row.get(inv["debit"])) if "debit" in inv else 0.0
            credit = _to_amount(raw_row.get(inv["credit"])) if "credit" in inv else 0.0
        except ValueError as exc:
            errors.append({"row": line_no, "reason": str(exc)})
            continue
        if debit == 0.0 and credit == 0.0:
            errors.append({"row": line_no, "reason": "débit et crédit tous deux à zéro"})
            continue
        label_raw = raw_row.get(inv["label"]) if "label" in inv else None
        label = str(label_raw).strip() if label_raw is not None else ""
        if label.lower() == "nan":
            label = ""
        rows.append({"account": account, "label": label or None, "debit": debit, "credit": credit})

    return rows, errors
